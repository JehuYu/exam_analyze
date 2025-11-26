#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
成绩分析核心模块
包含数据处理和统计计算的核心逻辑
"""

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
import math
import matplotlib.pyplot as plt
import matplotlib
from matplotlib import font_manager
import os

warnings.filterwarnings('ignore')

# 设置中文字体
matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS']
matplotlib.rcParams['axes.unicode_minus'] = False


class SubjectConfig:
    """学科配置类"""
    def __init__(self, name, max_score=150, pass_percent=60, excellence_percent=80):
        self.name = name
        self.max_score = max_score
        self.pass_percent = pass_percent
        self.excellence_percent = excellence_percent


class SubjectManager:
    """学科管理器"""
    
    def __init__(self):
        self.subjects = []
        
    def add_subject(self, config):
        if not any(s.name == config.name for s in self.subjects):
            self.subjects.append(config)
            return True
        return False
    
    def remove_subject(self, name):
        self.subjects = [s for s in self.subjects if s.name != name]
    
    def update_subject(self, name, config):
        for i, s in enumerate(self.subjects):
            if s.name == name:
                self.subjects[i] = config
                return True
        return False
    
    def get_subject(self, name):
        for s in self.subjects:
            if s.name == name:
                return s
        return None
    
    def get_subjects(self):
        return self.subjects.copy()
    
    def clear(self):
        self.subjects = []
    
    def auto_detect_from_excel(self, excel_file):
        """从Excel自动识别学科"""
        try:
            df = pd.read_excel(excel_file)
            cols = df.columns.tolist()
            
            subject_scores = {
                '语文': 150, '数学': 150, '外语': 120, '英语': 120,
                '科学': 180, '社会': 100, '社政': 100, '体育': 40,
            }
            
            detected = []
            i = 4
            
            while i < len(cols):
                if i + 1 < len(cols):
                    subject_col = cols[i + 1]
                    max_score = 150
                    for key, score in subject_scores.items():
                        if key in subject_col:
                            max_score = score
                            break
                    config = SubjectConfig(subject_col, max_score)
                    detected.append(config)
                    i += 2
                else:
                    break
            
            self.subjects = detected
            return True, detected
            
        except Exception as e:
            return False, str(e)


class GradeAnalysisCore:
    """成绩分析核心类"""
    
    def __init__(self, excel_file, subject_manager):
        self.excel_file = excel_file
        self.df = None
        self.subject_manager = subject_manager
        self.subjects = [s.name for s in subject_manager.get_subjects()]
        self.schools = []
        self.statistics = {}
        
    def load_data(self):
        """加载Excel数据"""
        try:
            self.df = pd.read_excel(self.excel_file)
            print(f"✓ 成功读取文件: {self.excel_file}")
            print(f"✓ 数据行数: {len(self.df)}")
        except Exception as e:
            print(f"✗ 读取文件失败: {e}")
            return False
        
        self._clean_data()
        return True
    
    def _clean_data(self):
        """数据清洗"""
        print("\n数据清洗中...")
        
        cleaned_data = {
            '考号': self.df['考号'],
            '姓名': self.df['姓名'],
            '学校名称': self.df['学校名称'],
            '班级名称': self.df['班级名称'],
        }
        
        cols = self.df.columns.tolist()
        
        for subject in self.subjects:
            found = False
            for i, col in enumerate(cols):
                if subject in col and i > 4:
                    if i > 0:
                        absent_col = cols[i - 1]
                        score_col = cols[i]
                        cleaned_data[f'{subject}_缺考'] = self.df[absent_col]
                        cleaned_data[f'{subject}'] = pd.to_numeric(self.df[score_col], errors='coerce')
                        found = True
                        break

            if not found:
                print(f"警告: 未找到科目 {subject}")
                cleaned_data[f'{subject}_缺考'] = 'N'
                cleaned_data[f'{subject}'] = 0

        if '备注' in self.df.columns:
            cleaned_data['备注'] = self.df['备注']

        self.df = pd.DataFrame(cleaned_data)
        self.schools = sorted(self.df['学校名称'].unique().tolist())
        print(f"✓ 数据清洗完成，学校数量: {len(self.schools)}")

    def calculate_statistics(self):
        """计算统计数据"""
        print("\n计算统计数据...")

        for subject in self.subjects:
            print(f"  计算 {subject}...")
            self.statistics[subject] = self._calculate_subject_stats(subject)

        print(f"  计算总分...")
        self.statistics['总分'] = self._calculate_total_stats()

        print("✓ 所有统计计算完成")

    def _calculate_subject_stats(self, subject):
        """计算单个科目的统计数据"""
        config = self.subject_manager.get_subject(subject)
        max_score = config.max_score
        pass_line = max_score * config.pass_percent / 100
        excellence_line = max_score * config.excellence_percent / 100

        stats = {}

        all_students = self.df[self.df[f'{subject}_缺考'] == 'N']
        all_scores = all_students[subject].dropna()

        total_count = len(all_scores)

        if total_count == 0:
            return {'data': pd.DataFrame(), 'pass_line': pass_line, 'excellence_line': excellence_line}

        sorted_scores = all_scores.sort_values(ascending=False).reset_index(drop=True)

        top30_count = math.floor(total_count * 0.3)
        if top30_count > 0 and top30_count <= total_count:
            top30_line = sorted_scores.iloc[top30_count - 1]
        else:
            top30_line = sorted_scores.max()

        top80_count = math.floor(total_count * 0.8)
        if top80_count < total_count:
            bottom20_line = sorted_scores.iloc[top80_count]
        else:
            bottom20_line = sorted_scores.min()

        school_stats = []

        for school in self.schools:
            school_data = self.df[self.df['学校名称'] == school]
            school_students = school_data[school_data[f'{subject}_缺考'] == 'N']
            school_scores = school_students[subject].dropna()

            if len(school_scores) == 0:
                continue

            count = len(school_scores)
            pass_count = (school_scores >= pass_line).sum()
            pass_rate = pass_count / count * 100
            excellence_count = (school_scores >= excellence_line).sum()
            excellence_rate = excellence_count / count * 100
            avg_score = school_scores.mean()
            bottom20_count = (school_scores <= bottom20_line).sum()
            bottom20_rate = bottom20_count / count * 100
            top30_count_school = (school_scores >= top30_line).sum()
            top30_rate = top30_count_school / count * 100

            school_stats.append({
                '学校': school,
                '考试人数': count,
                '合格率': pass_rate,
                '优秀率': excellence_rate,
                '平均分': avg_score,
                '后20%': bottom20_rate,
                '前30%': top30_rate,
            })

        df_stats = pd.DataFrame(school_stats)

        df_stats['合格率排序'] = df_stats['合格率'].rank(ascending=False, method='min').astype(int)
        df_stats['优秀率排序'] = df_stats['优秀率'].rank(ascending=False, method='min').astype(int)
        df_stats['平均分排序'] = df_stats['平均分'].rank(ascending=False, method='min').astype(int)
        df_stats['后20%排序'] = df_stats['后20%'].rank(ascending=True, method='min').astype(int)
        df_stats['前30%排序'] = df_stats['前30%'].rank(ascending=False, method='min').astype(int)

        all_pass_rate = (all_scores >= pass_line).sum() / total_count * 100
        all_excellence_rate = (all_scores >= excellence_line).sum() / total_count * 100
        all_avg = all_scores.mean()
        all_bottom20 = (all_scores <= bottom20_line).sum() / total_count * 100
        all_top30 = (all_scores >= top30_line).sum() / total_count * 100

        all_row = pd.DataFrame([{
            '学校': '全区',
            '考试人数': total_count,
            '合格率': all_pass_rate,
            '优秀率': all_excellence_rate,
            '平均分': all_avg,
            '后20%': all_bottom20,
            '前30%': all_top30,
            '合格率排序': '',
            '优秀率排序': '',
            '平均分排序': '',
            '后20%排序': '',
            '前30%排序': '',
        }])

        df_stats = pd.concat([df_stats, all_row], ignore_index=True)

        stats['data'] = df_stats
        stats['pass_line'] = pass_line
        stats['excellence_line'] = excellence_line
        stats['top30_line'] = top30_line
        stats['bottom20_line'] = bottom20_line
        stats['max_score'] = max_score

        return stats

    def _calculate_total_stats(self):
        """计算总分统计"""
        self.df['总分'] = 0
        self.df['有效总分'] = True

        total_max_score = sum(s.max_score for s in self.subject_manager.get_subjects())
        avg_pass_percent = sum(s.pass_percent for s in self.subject_manager.get_subjects()) / len(self.subject_manager.get_subjects())
        avg_excellence_percent = sum(s.excellence_percent for s in self.subject_manager.get_subjects()) / len(self.subject_manager.get_subjects())

        pass_line = total_max_score * avg_pass_percent / 100
        excellence_line = total_max_score * avg_excellence_percent / 100

        for idx, row in self.df.iterrows():
            total = 0
            valid = True

            for subject in self.subjects:
                if row[f'{subject}_缺考'] == 'Y' or pd.isna(row[subject]):
                    valid = False
                    break
                total += row[subject]

            self.df.at[idx, '总分'] = total if valid else 0
            self.df.at[idx, '有效总分'] = valid

        valid_students = self.df[self.df['有效总分'] == True]
        all_scores = valid_students['总分']

        total_count = len(all_scores)

        if total_count == 0:
            return {'data': pd.DataFrame(), 'pass_line': pass_line, 'excellence_line': excellence_line}

        sorted_scores = all_scores.sort_values(ascending=False).reset_index(drop=True)

        top30_count = math.floor(total_count * 0.3)
        if top30_count > 0:
            top30_line = sorted_scores.iloc[top30_count - 1]
        else:
            top30_line = sorted_scores.max()

        top80_count = math.floor(total_count * 0.8)
        if top80_count < total_count:
            bottom20_line = sorted_scores.iloc[top80_count]
        else:
            bottom20_line = sorted_scores.min()

        school_stats = []

        for school in self.schools:
            school_data = valid_students[valid_students['学校名称'] == school]
            school_scores = school_data['总分']

            if len(school_scores) == 0:
                continue

            count = len(school_scores)
            pass_count = (school_scores >= pass_line).sum()
            pass_rate = pass_count / count * 100
            excellence_count = (school_scores >= excellence_line).sum()
            excellence_rate = excellence_count / count * 100
            avg_score = school_scores.mean()
            bottom20_count = (school_scores <= bottom20_line).sum()
            bottom20_rate = bottom20_count / count * 100
            top30_count_school = (school_scores >= top30_line).sum()
            top30_rate = top30_count_school / count * 100

            school_stats.append({
                '学校': school,
                '考试人数': count,
                '合格率': pass_rate,
                '优秀率': excellence_rate,
                '平均分': avg_score,
                '后20%': bottom20_rate,
                '前30%': top30_rate,
            })

        df_stats = pd.DataFrame(school_stats)

        df_stats['合格率排序'] = df_stats['合格率'].rank(ascending=False, method='min').astype(int)
        df_stats['优秀率排序'] = df_stats['优秀率'].rank(ascending=False, method='min').astype(int)
        df_stats['平均分排序'] = df_stats['平均分'].rank(ascending=False, method='min').astype(int)
        df_stats['后20%排序'] = df_stats['后20%'].rank(ascending=True, method='min').astype(int)
        df_stats['前30%排序'] = df_stats['前30%'].rank(ascending=False, method='min').astype(int)

        all_pass_rate = (all_scores >= pass_line).sum() / total_count * 100
        all_excellence_rate = (all_scores >= excellence_line).sum() / total_count * 100
        all_avg = all_scores.mean()
        all_bottom20 = (all_scores <= bottom20_line).sum() / total_count * 100
        all_top30 = (all_scores >= top30_line).sum() / total_count * 100

        all_row = pd.DataFrame([{
            '学校': '全区',
            '考试人数': total_count,
            '合格率': all_pass_rate,
            '优秀率': all_excellence_rate,
            '平均分': all_avg,
            '后20%': all_bottom20,
            '前30%': all_top30,
            '合格率排序': '',
            '优秀率排序': '',
            '平均分排序': '',
            '后20%排序': '',
            '前30%排序': '',
        }])

        df_stats = pd.concat([df_stats, all_row], ignore_index=True)

        return {
            'data': df_stats,
            'pass_line': pass_line,
            'excellence_line': excellence_line,
            'top30_line': top30_line,
            'bottom20_line': bottom20_line,
            'max_score': total_max_score,
        }

    def generate_charts(self, output_dir='.'):
        """生成所有图表"""
        print("\n生成图表...")

        chart_files = {}

        # 1. 各科平均分对比图
        chart_files['avg_comparison'] = self._generate_avg_comparison_chart(output_dir)

        # 2. 各校综合表现雷达图
        chart_files['school_radar'] = self._generate_school_radar_chart(output_dir)

        # 3. 分数段分布图
        chart_files['score_distribution'] = self._generate_score_distribution_chart(output_dir)

        # 4. 合格率优秀率对比图
        chart_files['rate_comparison'] = self._generate_rate_comparison_chart(output_dir)

        # 5. 前30%后20%对比图
        chart_files['cutoff_comparison'] = self._generate_cutoff_comparison_chart(output_dir)

        print("✓ 所有图表生成完成")
        return chart_files

    def _generate_avg_comparison_chart(self, output_dir):
        """生成各科平均分对比柱状图"""
        fig, ax = plt.subplots(figsize=(14, 8))

        # 准备数据
        schools = [s for s in self.schools if s != '全区']
        x = np.arange(len(schools))
        width = 0.15

        # 为每个科目绘制柱状图
        for i, subject in enumerate(self.subjects):
            stats = self.statistics[subject]['data']
            stats = stats[stats['学校'] != '全区']

            avg_scores = []
            for school in schools:
                school_data = stats[stats['学校'] == school]
                if not school_data.empty:
                    avg_scores.append(school_data.iloc[0]['平均分'])
                else:
                    avg_scores.append(0)

            ax.bar(x + i * width, avg_scores, width, label=subject)

        ax.set_xlabel('学校', fontsize=12, fontweight='bold')
        ax.set_ylabel('平均分', fontsize=12, fontweight='bold')
        ax.set_title('各学校各科目平均分对比', fontsize=16, fontweight='bold', pad=20)
        ax.set_xticks(x + width * (len(self.subjects) - 1) / 2)
        ax.set_xticklabels(schools, rotation=45, ha='right')
        ax.legend(loc='upper left', fontsize=10)
        ax.grid(axis='y', alpha=0.3, linestyle='--')

        plt.tight_layout()
        output_file = os.path.join(output_dir, '各科平均分对比图.png')
        plt.savefig(output_file, dpi=300, bbox_inches='tight')
        plt.close()

        print(f"  ✓ 各科平均分对比图")
        return output_file

    def _generate_school_radar_chart(self, output_dir):
        """生成学校综合表现雷达图（前5所学校）"""
        fig, ax = plt.subplots(figsize=(12, 12), subplot_kw=dict(projection='polar'))

        # 选择前5所学校（按总分平均分排序）
        total_stats = self.statistics['总分']['data']
        total_stats = total_stats[total_stats['学校'] != '全区']
        top_schools = total_stats.nlargest(5, '平均分')['学校'].tolist()

        # 准备维度
        categories = ['合格率', '优秀率', '平均分', '前30%', '后20%']
        N = len(categories)
        angles = [n / float(N) * 2 * np.pi for n in range(N)]
        angles += angles[:1]

        # 为每所学校绘制雷达图
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#FFA07A', '#98D8C8']

        for idx, school in enumerate(top_schools):
            school_data = total_stats[total_stats['学校'] == school].iloc[0]

            # 归一化数据（转换为0-100的范围）
            values = [
                school_data['合格率'],
                school_data['优秀率'],
                school_data['平均分'] / self.statistics['总分']['max_score'] * 100,
                school_data['前30%'],
                100 - school_data['后20%']  # 后20%越低越好，所以取反
            ]
            values += values[:1]

            ax.plot(angles, values, 'o-', linewidth=2, label=school, color=colors[idx])
            ax.fill(angles, values, alpha=0.15, color=colors[idx])

        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(categories, fontsize=11)
        ax.set_ylim(0, 100)
        ax.set_title('学校综合表现雷达图（Top 5）', fontsize=16, fontweight='bold', pad=30)
        ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize=10)
        ax.grid(True, alpha=0.3)

        plt.tight_layout()
        output_file = os.path.join(output_dir, '学校综合表现雷达图.png')
        plt.savefig(output_file, dpi=300, bbox_inches='tight')
        plt.close()

        print(f"  ✓ 学校综合表现雷达图")
        return output_file

    def _generate_score_distribution_chart(self, output_dir):
        """生成分数段分布图"""
        fig, axes = plt.subplots(2, 3, figsize=(18, 12))
        fig.suptitle('各科目分数段分布情况', fontsize=18, fontweight='bold', y=0.995)

        axes = axes.flatten()

        # 为每个科目生成分数段分布图
        for idx, subject in enumerate(self.subjects):
            ax = axes[idx]

            # 获取该科目的所有成绩
            subject_data = self.df[self.df[f'{subject}_缺考'] == 'N']
            scores = subject_data[subject].dropna()

            max_score = self.statistics[subject]['max_score']

            # 定义分数段
            bins = [0, max_score*0.4, max_score*0.6, max_score*0.8, max_score*0.9, max_score]
            labels = ['0-40%', '40-60%', '60-80%', '80-90%', '90-100%']
            colors = ['#FF6B6B', '#FFA07A', '#FFD93D', '#6BCF7F', '#4ECDC4']

            # 统计各分数段人数
            counts, _ = np.histogram(scores, bins=bins)

            # 绘制柱状图
            x = np.arange(len(labels))
            bars = ax.bar(x, counts, color=colors, alpha=0.8, edgecolor='black', linewidth=1.5)

            # 添加数值标签
            for i, (bar, count) in enumerate(zip(bars, counts)):
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{int(count)}\n({count/len(scores)*100:.1f}%)',
                       ha='center', va='bottom', fontsize=9, fontweight='bold')

            ax.set_xlabel('分数段', fontsize=11, fontweight='bold')
            ax.set_ylabel('人数', fontsize=11, fontweight='bold')
            ax.set_title(f'{subject}（满分{int(max_score)}分）', fontsize=13, fontweight='bold', pad=10)
            ax.set_xticks(x)
            ax.set_xticklabels(labels, rotation=0)
            ax.grid(axis='y', alpha=0.3, linestyle='--')

        # 隐藏多余的子图
        if len(self.subjects) < 6:
            for idx in range(len(self.subjects), 6):
                axes[idx].axis('off')

        plt.tight_layout()
        output_file = os.path.join(output_dir, '分数段分布图.png')
        plt.savefig(output_file, dpi=300, bbox_inches='tight')
        plt.close()

        print(f"  ✓ 分数段分布图")
        return output_file

    def _generate_rate_comparison_chart(self, output_dir):
        """生成合格率优秀率对比图"""
        fig, axes = plt.subplots(1, 2, figsize=(16, 7))

        schools = [s for s in self.schools if s != '全区']
        x = np.arange(len(schools))
        width = 0.15

        # 左图：合格率对比
        ax1 = axes[0]
        for i, subject in enumerate(self.subjects):
            stats = self.statistics[subject]['data']
            stats = stats[stats['学校'] != '全区']

            rates = []
            for school in schools:
                school_data = stats[stats['学校'] == school]
                if not school_data.empty:
                    rates.append(school_data.iloc[0]['合格率'])
                else:
                    rates.append(0)

            ax1.bar(x + i * width, rates, width, label=subject)

        ax1.set_xlabel('学校', fontsize=12, fontweight='bold')
        ax1.set_ylabel('合格率 (%)', fontsize=12, fontweight='bold')
        ax1.set_title('各学校各科目合格率对比', fontsize=14, fontweight='bold', pad=15)
        ax1.set_xticks(x + width * (len(self.subjects) - 1) / 2)
        ax1.set_xticklabels(schools, rotation=45, ha='right')
        ax1.legend(loc='upper left', fontsize=9)
        ax1.grid(axis='y', alpha=0.3, linestyle='--')
        ax1.axhline(y=60, color='red', linestyle='--', linewidth=1, alpha=0.5, label='60%基准线')

        # 右图：优秀率对比
        ax2 = axes[1]
        for i, subject in enumerate(self.subjects):
            stats = self.statistics[subject]['data']
            stats = stats[stats['学校'] != '全区']

            rates = []
            for school in schools:
                school_data = stats[stats['学校'] == school]
                if not school_data.empty:
                    rates.append(school_data.iloc[0]['优秀率'])
                else:
                    rates.append(0)

            ax2.bar(x + i * width, rates, width, label=subject)

        ax2.set_xlabel('学校', fontsize=12, fontweight='bold')
        ax2.set_ylabel('优秀率 (%)', fontsize=12, fontweight='bold')
        ax2.set_title('各学校各科目优秀率对比', fontsize=14, fontweight='bold', pad=15)
        ax2.set_xticks(x + width * (len(self.subjects) - 1) / 2)
        ax2.set_xticklabels(schools, rotation=45, ha='right')
        ax2.legend(loc='upper left', fontsize=9)
        ax2.grid(axis='y', alpha=0.3, linestyle='--')

        plt.tight_layout()
        output_file = os.path.join(output_dir, '合格率优秀率对比图.png')
        plt.savefig(output_file, dpi=300, bbox_inches='tight')
        plt.close()

        print(f"  ✓ 合格率优秀率对比图")
        return output_file

    def _generate_cutoff_comparison_chart(self, output_dir):
        """生成前30%后20%对比图"""
        fig, ax = plt.subplots(figsize=(14, 8))

        schools = [s for s in self.schools if s != '全区']
        x = np.arange(len(schools))
        width = 0.35

        # 获取总分数据
        total_stats = self.statistics['总分']['data']
        total_stats = total_stats[total_stats['学校'] != '全区']

        top30_rates = []
        bottom20_rates = []

        for school in schools:
            school_data = total_stats[total_stats['学校'] == school]
            if not school_data.empty:
                top30_rates.append(school_data.iloc[0]['前30%'])
                bottom20_rates.append(school_data.iloc[0]['后20%'])
            else:
                top30_rates.append(0)
                bottom20_rates.append(0)

        bars1 = ax.bar(x - width/2, top30_rates, width, label='前30%', color='#4ECDC4', alpha=0.8)
        bars2 = ax.bar(x + width/2, bottom20_rates, width, label='后20%', color='#FF6B6B', alpha=0.8)

        # 添加数值标签
        for bars in [bars1, bars2]:
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{height:.1f}%',
                       ha='center', va='bottom', fontsize=9, fontweight='bold')

        ax.set_xlabel('学校', fontsize=12, fontweight='bold')
        ax.set_ylabel('比例 (%)', fontsize=12, fontweight='bold')
        ax.set_title('各学校前30%和后20%比例对比（总分）', fontsize=16, fontweight='bold', pad=20)
        ax.set_xticks(x)
        ax.set_xticklabels(schools, rotation=45, ha='right')
        ax.legend(fontsize=11)
        ax.grid(axis='y', alpha=0.3, linestyle='--')

        plt.tight_layout()
        output_file = os.path.join(output_dir, '前30后20对比图.png')
        plt.savefig(output_file, dpi=300, bbox_inches='tight')
        plt.close()

        print(f"  ✓ 前30%后20%对比图")
        return output_file

    def generate_analysis_report(self):
        """生成智能分析报告"""
        print("\n生成智能分析报告...")

        analysis = {
            'overall': self._analyze_overall(),
            'subjects': self._analyze_subjects(),
            'schools': self._analyze_schools(),
            'recommendations': self._generate_recommendations()
        }

        print("✓ 智能分析报告生成完成")
        return analysis

    def _analyze_overall(self):
        """整体分析"""
        total_stats = self.statistics['总分']['data']
        all_data = total_stats[total_stats['学校'] == '全区'].iloc[0]

        analysis = {
            'summary': f"本次考试共有{int(all_data['考试人数'])}名学生参加，"
                      f"总分平均分为{all_data['平均分']:.2f}分，"
                      f"合格率{all_data['合格率']:.2f}%，"
                      f"优秀率{all_data['优秀率']:.2f}%。",
            'level': self._get_performance_level(all_data['合格率'], all_data['优秀率']),
            'pass_rate': all_data['合格率'],
            'excellence_rate': all_data['优秀率'],
            'avg_score': all_data['平均分']
        }

        return analysis

    def _analyze_subjects(self):
        """科目分析"""
        subject_analysis = []

        for subject in self.subjects:
            stats = self.statistics[subject]['data']
            all_data = stats[stats['学校'] == '全区'].iloc[0]

            # 找出表现最好和最差的学校
            school_stats = stats[stats['学校'] != '全区']
            best_school = school_stats.nlargest(1, '平均分').iloc[0]
            worst_school = school_stats.nsmallest(1, '平均分').iloc[0]

            analysis = {
                'subject': subject,
                'avg_score': all_data['平均分'],
                'pass_rate': all_data['合格率'],
                'excellence_rate': all_data['优秀率'],
                'best_school': best_school['学校'],
                'best_avg': best_school['平均分'],
                'worst_school': worst_school['学校'],
                'worst_avg': worst_school['平均分'],
                'gap': best_school['平均分'] - worst_school['平均分'],
                'summary': f"{subject}平均分{all_data['平均分']:.2f}分，"
                          f"合格率{all_data['合格率']:.2f}%，"
                          f"优秀率{all_data['优秀率']:.2f}%。"
                          f"{best_school['学校']}表现最好（{best_school['平均分']:.2f}分），"
                          f"{worst_school['学校']}需要加强（{worst_school['平均分']:.2f}分），"
                          f"校际差距{best_school['平均分'] - worst_school['平均分']:.2f}分。"
            }

            subject_analysis.append(analysis)

        return subject_analysis

    def _analyze_schools(self):
        """学校分析"""
        total_stats = self.statistics['总分']['data']
        school_stats = total_stats[total_stats['学校'] != '全区']

        school_analysis = []

        for idx, row in school_stats.iterrows():
            school = row['学校']

            # 计算该学校在各科的排名
            subject_ranks = []
            for subject in self.subjects:
                stats = self.statistics[subject]['data']
                school_data = stats[stats['学校'] == school]
                if not school_data.empty:
                    rank = school_data.iloc[0]['平均分排序']
                    subject_ranks.append(rank)

            avg_rank = np.mean(subject_ranks) if subject_ranks else 0

            # 找出优势科目和薄弱科目
            subject_avgs = []
            for subject in self.subjects:
                stats = self.statistics[subject]['data']
                school_data = stats[stats['学校'] == school]
                if not school_data.empty:
                    subject_avgs.append({
                        'subject': subject,
                        'avg': school_data.iloc[0]['平均分'],
                        'rank': school_data.iloc[0]['平均分排序']
                    })

            subject_avgs.sort(key=lambda x: x['rank'])

            strength = subject_avgs[0] if subject_avgs else None
            weakness = subject_avgs[-1] if subject_avgs else None

            analysis = {
                'school': school,
                'total_avg': row['平均分'],
                'total_rank': row['平均分排序'],
                'pass_rate': row['合格率'],
                'excellence_rate': row['优秀率'],
                'top30_rate': row['前30%'],
                'bottom20_rate': row['后20%'],
                'avg_rank': avg_rank,
                'strength_subject': strength['subject'] if strength else None,
                'strength_avg': strength['avg'] if strength else None,
                'weakness_subject': weakness['subject'] if weakness else None,
                'weakness_avg': weakness['avg'] if weakness else None,
                'summary': f"{school}总分平均{row['平均分']:.2f}分，排名第{int(row['平均分排序'])}，"
                          f"合格率{row['合格率']:.2f}%，优秀率{row['优秀率']:.2f}%。"
                          f"优势科目为{strength['subject'] if strength else '无'}，"
                          f"薄弱科目为{weakness['subject'] if weakness else '无'}。"
            }

            school_analysis.append(analysis)

        return school_analysis

    def _generate_recommendations(self):
        """生成改进建议"""
        recommendations = []

        # 基于整体数据的建议
        total_stats = self.statistics['总分']['data']
        all_data = total_stats[total_stats['学校'] == '全区'].iloc[0]

        if all_data['合格率'] < 70:
            recommendations.append({
                'type': '整体建议',
                'level': '重要',
                'content': f"全区合格率仅{all_data['合格率']:.2f}%，低于70%的理想水平，建议加强基础知识教学，提高整体合格率。"
            })

        if all_data['优秀率'] < 10:
            recommendations.append({
                'type': '整体建议',
                'level': '重要',
                'content': f"全区优秀率仅{all_data['优秀率']:.2f}%，建议加强拔尖学生培养，提升优秀率。"
            })

        # 基于科目的建议
        for subject in self.subjects:
            stats = self.statistics[subject]['data']
            all_data = stats[stats['学校'] == '全区'].iloc[0]

            if all_data['合格率'] < 60:
                recommendations.append({
                    'type': '科目建议',
                    'level': '紧急',
                    'content': f"{subject}合格率仅{all_data['合格率']:.2f}%，需要重点关注，建议分析试卷难度和教学方法。"
                })

        # 基于学校的建议
        school_stats = total_stats[total_stats['学校'] != '全区']
        bottom_schools = school_stats.nsmallest(3, '平均分')

        for idx, row in bottom_schools.iterrows():
            if row['后20%'] > 30:
                recommendations.append({
                    'type': '学校建议',
                    'level': '关注',
                    'content': f"{row['学校']}后20%比例达{row['后20%']:.2f}%，建议加强后进生辅导，减少两极分化。"
                })

        return recommendations

    def _get_performance_level(self, pass_rate, excellence_rate):
        """评估整体表现水平"""
        if pass_rate >= 85 and excellence_rate >= 20:
            return '优秀'
        elif pass_rate >= 75 and excellence_rate >= 10:
            return '良好'
        elif pass_rate >= 60 and excellence_rate >= 5:
            return '中等'
        else:
            return '需要改进'

    def generate_word_report(self, output_file, progress_callback=None):
        """生成Word统计报告（含图表和分析）"""
        print("\n生成Word统计报告...")

        doc = Document()
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(10.5)

        total_steps = len(self.subjects) + 8
        current_step = 0

        # 获取输出目录
        output_dir = os.path.dirname(output_file) or '.'

        # 添加封面
        if progress_callback:
            current_step += 1
            progress_callback(current_step / total_steps, "生成封面...")
        self._add_cover_page(doc)

        # 添加整体分析
        if progress_callback:
            current_step += 1
            progress_callback(current_step / total_steps, "生成整体分析...")
        analysis = self.generate_analysis_report()
        self._add_overall_analysis(doc, analysis)

        # 生成图表
        if progress_callback:
            current_step += 1
            progress_callback(current_step / total_steps, "生成图表...")
        chart_files = self.generate_charts(output_dir)

        # 添加图表到文档
        if progress_callback:
            current_step += 1
            progress_callback(current_step / total_steps, "插入图表...")
        self._add_charts_to_doc(doc, chart_files)

        # 添加各科统计表
        for subject in self.subjects:
            if progress_callback:
                current_step += 1
                progress_callback(current_step / total_steps, f"生成{subject}统计表...")
            self._add_subject_table(doc, subject)

        # 添加总分统计表
        if progress_callback:
            current_step += 1
            progress_callback(current_step / total_steps, "生成总分统计表...")
        self._add_total_table(doc)

        # 添加分数段分布表
        if progress_callback:
            current_step += 1
            progress_callback(current_step / total_steps, "生成分数段分布表...")
        self._add_score_distribution_table(doc)

        # 添加分数线表
        if progress_callback:
            current_step += 1
            progress_callback(current_step / total_steps, "生成分数线表...")
        self._add_cutoff_table(doc)

        # 添加详细分析报告
        self._add_detailed_analysis(doc, analysis)

        doc.save(output_file)
        print(f"✓ Word报告已生成: {output_file}")

        if progress_callback:
            progress_callback(1.0, "报告生成完成！")

        return output_file

    def _add_cover_page(self, doc):
        """添加封面"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('成绩统计分析报告')
        run.font.size = Pt(26)
        run.font.bold = True

        doc.add_paragraph()
        doc.add_paragraph()

        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(f'考试科目：{" ".join(self.subjects)}\n')
        run.font.size = Pt(14)

        run = info.add_run(f'参考学校：{len(self.schools)}所\n')
        run.font.size = Pt(14)

        total_students = self.statistics['总分']['data']
        all_data = total_students[total_students['学校'] == '全区'].iloc[0]
        run = info.add_run(f'参考人数：{int(all_data["考试人数"])}人\n')
        run.font.size = Pt(14)

        doc.add_page_break()

    def _add_overall_analysis(self, doc, analysis):
        """添加整体分析"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('一、整体分析')
        run.font.size = Pt(16)
        run.font.bold = True

        doc.add_paragraph()

        # 整体概况
        p = doc.add_paragraph()
        run = p.add_run('1. 整体概况\n')
        run.font.size = Pt(14)
        run.font.bold = True

        p = doc.add_paragraph(analysis['overall']['summary'])
        p.paragraph_format.first_line_indent = Pt(21)

        p = doc.add_paragraph(f"整体表现水平：{analysis['overall']['level']}")
        p.paragraph_format.first_line_indent = Pt(21)

        doc.add_paragraph()

    def _add_charts_to_doc(self, doc, chart_files):
        """将图表添加到文档"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('二、数据可视化分析')
        run.font.size = Pt(16)
        run.font.bold = True

        doc.add_paragraph()

        # 添加各个图表
        chart_titles = {
            'avg_comparison': '1. 各学校各科目平均分对比',
            'school_radar': '2. 学校综合表现雷达图',
            'score_distribution': '3. 各科目分数段分布情况',
            'rate_comparison': '4. 合格率优秀率对比',
            'cutoff_comparison': '5. 前30%后20%比例对比'
        }

        for key, title_text in chart_titles.items():
            if key in chart_files and os.path.exists(chart_files[key]):
                p = doc.add_paragraph()
                run = p.add_run(title_text)
                run.font.size = Pt(12)
                run.font.bold = True

                doc.add_picture(chart_files[key], width=Inches(6.5))
                doc.add_paragraph()

        doc.add_page_break()

    def _add_detailed_analysis(self, doc, analysis):
        """添加详细分析报告"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('三、详细分析与建议')
        run.font.size = Pt(16)
        run.font.bold = True

        doc.add_paragraph()

        # 科目分析
        p = doc.add_paragraph()
        run = p.add_run('（一）科目分析\n')
        run.font.size = Pt(14)
        run.font.bold = True

        for idx, subj_analysis in enumerate(analysis['subjects'], 1):
            p = doc.add_paragraph(f"{idx}. {subj_analysis['summary']}")
            p.paragraph_format.first_line_indent = Pt(21)

        doc.add_paragraph()

        # 学校分析（Top 5）
        p = doc.add_paragraph()
        run = p.add_run('（二）学校分析（Top 5）\n')
        run.font.size = Pt(14)
        run.font.bold = True

        school_analysis = sorted(analysis['schools'], key=lambda x: x['total_rank'])[:5]

        for idx, school_data in enumerate(school_analysis, 1):
            p = doc.add_paragraph(f"{idx}. {school_data['summary']}")
            p.paragraph_format.first_line_indent = Pt(21)

        doc.add_paragraph()

        # 改进建议
        p = doc.add_paragraph()
        run = p.add_run('（三）改进建议\n')
        run.font.size = Pt(14)
        run.font.bold = True

        if analysis['recommendations']:
            for idx, rec in enumerate(analysis['recommendations'], 1):
                p = doc.add_paragraph(f"{idx}. 【{rec['level']}】{rec['content']}")
                p.paragraph_format.first_line_indent = Pt(21)
        else:
            p = doc.add_paragraph("整体表现良好，继续保持！")
            p.paragraph_format.first_line_indent = Pt(21)

    def export_to_excel(self, output_file):
        """导出详细数据到Excel"""
        print("\n导出Excel文件...")

        with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
            # 导出各科目统计
            for subject in self.subjects:
                stats = self.statistics[subject]['data']
                stats.to_excel(writer, sheet_name=subject, index=False)

            # 导出总分统计
            total_stats = self.statistics['总分']['data']
            total_stats.to_excel(writer, sheet_name='总分', index=False)

            # 导出原始数据（带总分）
            export_df = self.df.copy()
            export_df.to_excel(writer, sheet_name='原始数据', index=False)

            # 导出分析报告
            analysis = self.generate_analysis_report()

            # 整体分析
            overall_df = pd.DataFrame([{
                '项目': '整体概况',
                '内容': analysis['overall']['summary'],
                '表现水平': analysis['overall']['level']
            }])
            overall_df.to_excel(writer, sheet_name='整体分析', index=False)

            # 科目分析
            subject_analysis_data = []
            for subj in analysis['subjects']:
                subject_analysis_data.append({
                    '科目': subj['subject'],
                    '平均分': subj['avg_score'],
                    '合格率': subj['pass_rate'],
                    '优秀率': subj['excellence_rate'],
                    '最好学校': subj['best_school'],
                    '最好平均分': subj['best_avg'],
                    '最差学校': subj['worst_school'],
                    '最差平均分': subj['worst_avg'],
                    '校际差距': subj['gap']
                })

            subject_analysis_df = pd.DataFrame(subject_analysis_data)
            subject_analysis_df.to_excel(writer, sheet_name='科目分析', index=False)

            # 学校分析
            school_analysis_data = []
            for school in analysis['schools']:
                school_analysis_data.append({
                    '学校': school['school'],
                    '总分平均': school['total_avg'],
                    '总分排名': school['total_rank'],
                    '合格率': school['pass_rate'],
                    '优秀率': school['excellence_rate'],
                    '前30%': school['top30_rate'],
                    '后20%': school['bottom20_rate'],
                    '优势科目': school['strength_subject'],
                    '薄弱科目': school['weakness_subject']
                })

            school_analysis_df = pd.DataFrame(school_analysis_data)
            school_analysis_df.to_excel(writer, sheet_name='学校分析', index=False)

            # 改进建议
            if analysis['recommendations']:
                rec_data = []
                for rec in analysis['recommendations']:
                    rec_data.append({
                        '类型': rec['type'],
                        '级别': rec['level'],
                        '建议内容': rec['content']
                    })

                rec_df = pd.DataFrame(rec_data)
                rec_df.to_excel(writer, sheet_name='改进建议', index=False)

        print(f"✓ Excel文件已导出: {output_file}")
        return output_file

    def _add_subject_table(self, doc, subject):
        """添加单个科目的统计表"""
        stats = self.statistics[subject]
        df_stats = stats['data']

        if len(df_stats) == 0:
            return

        max_score = stats['max_score']
        pass_line = stats['pass_line']
        excellence_line = stats['excellence_line']

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f'成绩统计表')
        run.font.size = Pt(12)
        run.font.bold = True

        table = doc.add_table(rows=len(df_stats) + 2, cols=12)
        table.style = 'Table Grid'
        self._set_table_border(table)

        row0 = table.rows[0]
        cell = row0.cells[0]
        for i in range(1, 12):
            cell.merge(row0.cells[i])
        cell.text = f'{subject}（{int(max_score)}分）'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)

        headers = ['学校', '考试\n人数', f'合格率\n(%-{int(pass_line)})', '排序',
                   f'优秀率\n(%-{int(excellence_line)})', '排序', '平均分', '排序',
                   '后20％', '排序', '前30％', '排序']

        row1 = table.rows[1]
        for i, header in enumerate(headers):
            cell = row1.cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)

        for idx, row_data in df_stats.iterrows():
            row = table.rows[idx + 2]

            values = [
                row_data['学校'],
                str(int(row_data['考试人数'])),
                f"{row_data['合格率']:.2f}%",
                str(row_data['合格率排序']) if row_data['合格率排序'] != '' else '',
                f"{row_data['优秀率']:.2f}%",
                str(row_data['优秀率排序']) if row_data['优秀率排序'] != '' else '',
                f"{row_data['平均分']:.2f}",
                str(row_data['平均分排序']) if row_data['平均分排序'] != '' else '',
                f"{row_data['后20%']:.2f}%",
                str(row_data['后20%排序']) if row_data['后20%排序'] != '' else '',
                f"{row_data['前30%']:.2f}%",
                str(row_data['前30%排序']) if row_data['前30%排序'] != '' else '',
            ]

            for i, value in enumerate(values):
                cell = row.cells[i]
                cell.text = value
                if i > 0:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(10)

                if row_data['学校'] == '全区':
                    cell.paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

    def _add_total_table(self, doc):
        """添加总分统计表"""
        stats = self.statistics['总分']
        df_stats = stats['data']

        if len(df_stats) == 0:
            return

        max_score = stats['max_score']
        pass_line = stats['pass_line']
        excellence_line = stats['excellence_line']

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f'成绩统计表')
        run.font.size = Pt(12)
        run.font.bold = True

        table = doc.add_table(rows=len(df_stats) + 2, cols=12)
        table.style = 'Table Grid'
        self._set_table_border(table)

        row0 = table.rows[0]
        cell = row0.cells[0]
        for i in range(1, 12):
            cell.merge(row0.cells[i])
        cell.text = f'总分（{int(max_score)}分）'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)

        headers = ['学校', '考试\n人数', f'合格率\n(%-{int(pass_line)})', '排序',
                   f'优秀率\n(%-{int(excellence_line)})', '排序', '平均分', '排序',
                   '后20％', '排序', '前30％', '排序']

        row1 = table.rows[1]
        for i, header in enumerate(headers):
            cell = row1.cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)

        for idx, row_data in df_stats.iterrows():
            row = table.rows[idx + 2]

            values = [
                row_data['学校'],
                str(int(row_data['考试人数'])),
                f"{row_data['合格率']:.2f}%",
                str(row_data['合格率排序']) if row_data['合格率排序'] != '' else '',
                f"{row_data['优秀率']:.2f}%",
                str(row_data['优秀率排序']) if row_data['优秀率排序'] != '' else '',
                f"{row_data['平均分']:.2f}",
                str(row_data['平均分排序']) if row_data['平均分排序'] != '' else '',
                f"{row_data['后20%']:.2f}%",
                str(row_data['后20%排序']) if row_data['后20%排序'] != '' else '',
                f"{row_data['前30%']:.2f}%",
                str(row_data['前30%排序']) if row_data['前30%排序'] != '' else '',
            ]

            for i, value in enumerate(values):
                cell = row.cells[i]
                cell.text = value
                if i > 0:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(10)

                if row_data['学校'] == '全区':
                    cell.paragraphs[0].runs[0].font.bold = True

        note = doc.add_paragraph('注：计算总分时未参加所有科目考试的学生未统计在内')
        note.alignment = WD_ALIGN_PARAGRAPH.LEFT
        note.runs[0].font.size = Pt(9)

        doc.add_paragraph()

    def _add_score_distribution_table(self, doc):
        """添加分数段分布表"""
        valid_students = self.df[self.df['有效总分'] == True]

        if len(valid_students) == 0:
            return

        total_max = self.statistics['总分']['max_score']

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('期末考试总分分数段分布情况')
        run.font.size = Pt(12)
        run.font.bold = True

        segments = []
        step = 50
        for i in range(0, int(total_max), step):
            segments.append((i, i + step - 1))
        segments.append((int(total_max), int(total_max)))

        school_distribution = []

        for school in self.schools:
            school_data = valid_students[valid_students['学校名称'] == school]
            scores = school_data['总分']

            if len(scores) == 0:
                continue

            dist = {'学校': school}
            for start, end in segments:
                if start == end:
                    count = (scores == start).sum()
                    dist[f'{int(start)}'] = int(count)
                else:
                    count = ((scores >= start) & (scores < end + 1)).sum()
                    dist[f'{int(start)}-{int(end)}'] = int(count)

            school_distribution.append(dist)

        all_scores = valid_students['总分']
        all_dist = {'学校': '全区'}
        for start, end in segments:
            if start == end:
                count = (all_scores == start).sum()
                all_dist[f'{int(start)}'] = int(count)
            else:
                count = ((all_scores >= start) & (all_scores < end + 1)).sum()
                all_dist[f'{int(start)}-{int(end)}'] = int(count)

        school_distribution.append(all_dist)

        df_dist = pd.DataFrame(school_distribution)

        table = doc.add_table(rows=len(df_dist) + 1, cols=len(df_dist.columns))
        table.style = 'Table Grid'
        self._set_table_border(table)

        row0 = table.rows[0]
        for i, col in enumerate(df_dist.columns):
            cell = row0.cells[i]
            cell.text = col
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)

        for idx, row_data in df_dist.iterrows():
            row = table.rows[idx + 1]
            for i, col in enumerate(df_dist.columns):
                cell = row.cells[i]
                value = row_data[col]
                if col == '学校':
                    cell.text = str(value)
                else:
                    cell.text = str(int(value)) if pd.notna(value) else '0'
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(10)

                if row_data['学校'] == '全区':
                    cell.paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

    def _add_cutoff_table(self, doc):
        """添加前30%后20%分数线表"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('期末考试各科后20%、前30%分数线')
        run.font.size = Pt(12)
        run.font.bold = True

        table = doc.add_table(rows=3, cols=len(self.subjects) + 2)
        table.style = 'Table Grid'
        self._set_table_border(table)

        row0 = table.rows[0]
        row0.cells[0].text = '学科'
        row0.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row0.cells[0].paragraphs[0].runs[0].font.bold = True
        row0.cells[0].paragraphs[0].runs[0].font.size = Pt(10)

        for i, subject in enumerate(self.subjects):
            cell = row0.cells[i + 1]
            cell.text = subject
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)

        row0.cells[len(self.subjects) + 1].text = '总分'
        row0.cells[len(self.subjects) + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row0.cells[len(self.subjects) + 1].paragraphs[0].runs[0].font.bold = True
        row0.cells[len(self.subjects) + 1].paragraphs[0].runs[0].font.size = Pt(10)

        row1 = table.rows[1]
        row1.cells[0].text = '后20%'
        row1.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row1.cells[0].paragraphs[0].runs[0].font.bold = True
        row1.cells[0].paragraphs[0].runs[0].font.size = Pt(10)

        for i, subject in enumerate(self.subjects):
            cell = row1.cells[i + 1]
            bottom20_line = self.statistics[subject]['bottom20_line']
            cell.text = f'{bottom20_line:.1f}'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(10)

        cell = row1.cells[len(self.subjects) + 1]
        bottom20_line = self.statistics['总分']['bottom20_line']
        cell.text = f'{bottom20_line:.1f}'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)

        row2 = table.rows[2]
        row2.cells[0].text = '前30%'
        row2.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row2.cells[0].paragraphs[0].runs[0].font.bold = True
        row2.cells[0].paragraphs[0].runs[0].font.size = Pt(10)

        for i, subject in enumerate(self.subjects):
            cell = row2.cells[i + 1]
            top30_line = self.statistics[subject]['top30_line']
            cell.text = f'{top30_line:.1f}'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(10)

        cell = row2.cells[len(self.subjects) + 1]
        top30_line = self.statistics['总分']['top30_line']
        cell.text = f'{top30_line:.1f}'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    def _set_table_border(self, table):
        """设置表格边框"""
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)

        tblPr.append(tblBorders)

